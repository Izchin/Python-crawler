# 在网易云音乐打开某歌手主页 ，把网址上的数字复制并粘贴到以下程序的id输入
# 同时也可以搜索某歌单所有歌曲，输入歌单id即可
# 在程序所在目录下会创建相关歌手歌曲列表word文档

import re
import requests
from docx import Document
import urllib.request
from docx.shared import Inches

search = int(input(print("How do you want to search the playlist? by singer(1) or by playlist(2) ?")))
if search == 1:
    searchStyle = 'artist'
    Id = str(input(print('Please enter the id of the singer:')))
else:
    searchStyle = 'playlist'
    Id = str(input(print('Please enter the id of the playlist:')))
url = "https://music.163.com/" + searchStyle + "?id=" + Id
music_text = requests.get(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) '
                                                       'AppleWebKit/537.36 (KHTML, like Gecko) '
                                                       'Chrome/51.0.2704.63 Safari/537.36'}).text

music_code = re.findall(r'<ul class="f-hide">.*?</a></li></ul>', music_text, re.I)[0]
musicUrls = re.findall(r'<li><a href="(.*?)">', music_code)
musicNames = re.findall(r'>(.*?)</a></li><li><a href="/song\?id=', music_code)
image_url = re.findall(r'"images": \["(.*?)"\],', music_text)
image_f = open((Id+'.jpg'), 'wb')
request = urllib.request.urlopen(image_url[0])
image = request.read()
image_f.write(image)
image_f.close()
listNum = 1
document = Document()
document.add_heading('Music List', 0)
document.add_paragraph('This is the playlist you want:\n')
document.add_picture((Id+'.jpg'), width=Inches(2.25))
for musicUrl, musicName in zip(musicUrls, musicNames):
    music_Url = "https://music.163.com/#" + musicUrl
    if listNum != 1:
        words = str(listNum-1) + ') '+'[' + musicName + ']' + ' Url :  ' + music_Url
        document.add_paragraph(words)
    listNum += 1
document.add_page_break()
document.save('Id'+Id+'.docx')
